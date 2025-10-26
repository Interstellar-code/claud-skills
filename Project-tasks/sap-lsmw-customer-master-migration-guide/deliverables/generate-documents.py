#!/usr/bin/env python3
"""
Generate Word document and Excel template for SAP LSMW Customer Master Migration Guide
Uses python-docx for Word and openpyxl for Excel
"""

import os
import sys

# Word Document Generation
def generate_word_document():
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        print("ERROR: python-docx not installed. Install with: pip install python-docx")
        return False

    print("Generating Word document...")

    # Read markdown file
    with open('SAP-LSMW-Customer-Master-Migration-Guide.md', 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading('SAP LSMW Customer Master Migration Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    metadata = doc.add_paragraph()
    metadata.add_run('Version: ').bold = True
    metadata.add_run('1.0\n')
    metadata.add_run('Date: ').bold = True
    metadata.add_run('October 2025\n')
    metadata.add_run('Author: ').bold = True
    metadata.add_run('SAP Techno-Functional Consultant\n')
    metadata.add_run('Audience: ').bold = True
    metadata.add_run('SAP Consultants, Migration Teams\n')
    metadata.add_run('SAP Module: ').bold = True
    metadata.add_run('SD (Sales and Distribution)\n')
    metadata.add_run('Transaction: ').bold = True
    metadata.add_run('LSMW, XD01')
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        "1. Overview",
        "2. Customer Master Data Structure",
        "3. The 14-Step LSMW Process",
        "4. Recording Instructions",
        "5. Data File Preparation",
        "6. Testing Process",
        "7. Production Migration",
        "8. Troubleshooting",
        "9. Best Practices",
        "10. Appendices"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # Parse markdown content and add to document
    lines = content.split('\n')
    in_table = False
    table_data = []
    in_code_block = False
    code_block_lines = []

    for line in lines:
        # Skip initial metadata and TOC
        if line.startswith('**Version**') or line.startswith('**Date**') or line.startswith('**Author**'):
            continue
        if line.startswith('## Table of Contents'):
            # Skip the markdown TOC, we already added it
            in_toc = True
            continue

        # Handle headings
        if line.startswith('# ') and not line.startswith('##'):
            continue  # Skip title, already added
        elif line.startswith('## '):
            doc.add_heading(line[3:], 1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], 3)

        # Handle code blocks
        elif line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
            else:
                # End of code block
                in_code_block = False
                if code_block_lines:
                    p = doc.add_paragraph('\n'.join(code_block_lines))
                    p.style = 'Normal'
                    # Make code monospace and gray background
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
        elif in_code_block:
            code_block_lines.append(line)

        # Handle tables
        elif '|' in line and not in_code_block:
            if not in_table:
                in_table = True
                table_data = []
            table_data.append([cell.strip() for cell in line.split('|')[1:-1]])
        else:
            # End of table
            if in_table and table_data:
                # Remove separator row (usually second row with ---)
                if len(table_data) > 1 and all('---' in cell or ':-' in cell for cell in table_data[1]):
                    table_data.pop(1)

                if len(table_data) > 1:
                    # Create table
                    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                    table.style = 'Light Grid Accent 1'

                    # Populate table
                    for i, row_data in enumerate(table_data):
                        row = table.rows[i]
                        for j, cell_text in enumerate(row_data):
                            if j < len(row.cells):
                                row.cells[j].text = cell_text
                                # Bold header row
                                if i == 0:
                                    for paragraph in row.cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True

                in_table = False
                table_data = []

            # Handle regular paragraphs
            if line.strip() and not line.startswith('---'):
                # Handle bold and italic
                if '**' in line or '*' in line:
                    p = doc.add_paragraph()
                    # Simple bold handling
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # Odd indices are bold
                            run.bold = True
                else:
                    doc.add_paragraph(line)
            elif line.strip() == '---':
                pass  # Skip horizontal rules

    # Save document
    doc.save('SAP-LSMW-Customer-Master-Migration-Guide.docx')
    print("[OK] Word document created: SAP-LSMW-Customer-Master-Migration-Guide.docx")
    return True


# Excel Template Generation
def generate_excel_template():
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.worksheet.datavalidation import DataValidation
    except ImportError:
        print("ERROR: openpyxl not installed. Install with: pip install openpyxl")
        return False

    print("Generating Excel template...")

    wb = Workbook()

    # Define styles
    header_fill = PatternFill(start_color="0070C0", end_color="0070C0", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_alignment = Alignment(horizontal="center", vertical="center")

    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )

    # Sheet 1: General Data (KNA1)
    ws1 = wb.active
    ws1.title = "General Data (KNA1)"

    headers1 = ["KUNNR", "NAME1", "NAME2", "SORTL", "STRAS", "ORT01", "PSTLZ", "LAND1", "REGIO", "TELF1", "SMTP_ADDR", "KTOKD"]
    for col, header in enumerate(headers1, 1):
        cell = ws1.cell(1, col, header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Add data validation for Country (LAND1)
    dv_country = DataValidation(type="list", formula1='"US,GB,DE,FR,CN,JP,IN"', allow_blank=True)
    dv_country.error = 'Please select from dropdown'
    dv_country.errorTitle = 'Invalid Country Code'
    ws1.add_data_validation(dv_country)
    dv_country.add(f"H2:H1000")  # Column H = LAND1

    # Add data validation for Account Group (KTOKD)
    dv_ktokd = DataValidation(type="list", formula1='"KUNA,KUNB,KUNC"', allow_blank=True)
    dv_ktokd.error = 'Please select from dropdown'
    dv_ktokd.errorTitle = 'Invalid Account Group'
    ws1.add_data_validation(dv_ktokd)
    dv_ktokd.add(f"L2:L1000")  # Column L = KTOKD

    # Sample data (5 records)
    sample_data1 = [
        ["0000100001", "ACME Corporation", "", "ACME", "123 Main St", "New York", "10001", "US", "NY", "212-555-0100", "contact@acme.com", "KUNA"],
        ["0000100002", "Global Industries Ltd", "", "GLOBAL", "456 Park Ave", "London", "SW1A1AA", "GB", "", "+44-20-7123-4567", "info@global.com", "KUNA"],
        ["0000100003", "Tech Solutions GmbH", "", "TECH", "789 Hauptstr", "Berlin", "10115", "DE", "", "+49-30-1234567", "sales@tech.de", "KUNB"],
        ["0000100004", "Pacific Trade Co", "", "PACIFIC", "321 Ocean Blvd", "Tokyo", "100-0001", "JP", "", "+81-3-1234-5678", "contact@pacific.jp", "KUNA"],
        ["0000100005", "Delta Manufacturing", "", "DELTA", "654 Industrial Rd", "Paris", "75001", "FR", "", "+33-1-2345-6789", "info@delta.fr", "KUNC"]
    ]

    for row_data in sample_data1:
        ws1.append(row_data)

    # Sheet 2: Company Code Data (KNB1)
    ws2 = wb.create_sheet("Company Code (KNB1)")
    headers2 = ["KUNNR", "BUKRS", "AKONT", "ZTERM", "ZWELS"]
    for col, header in enumerate(headers2, 1):
        cell = ws2.cell(1, col, header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    sample_data2 = [
        ["0000100001", "1000", "140000", "Z030", "C"],
        ["0000100002", "1000", "140000", "Z030", "C"],
        ["0000100003", "2000", "140000", "Z060", "C"],
        ["0000100004", "1000", "140000", "Z030", "T"],
        ["0000100005", "2000", "140000", "Z045", "C"]
    ]

    for row_data in sample_data2:
        ws2.append(row_data)

    # Sheet 3: Sales Area Data (KNVV)
    ws3 = wb.create_sheet("Sales Area (KNVV)")
    headers3 = ["KUNNR", "VKORG", "VTWEG", "SPART", "KDGRP", "WAERS", "INCO1"]
    for col, header in enumerate(headers3, 1):
        cell = ws3.cell(1, col, header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment
        cell.border = thin_border

    # Add validation for Currency
    dv_waers = DataValidation(type="list", formula1='"USD,EUR,GBP,JPY,CNY"', allow_blank=True)
    dv_waers.error = 'Please select from dropdown'
    dv_waers.errorTitle = 'Invalid Currency'
    ws3.add_data_validation(dv_waers)
    dv_waers.add(f"F2:F1000")  # Column F = WAERS

    # Add validation for Incoterms
    dv_inco = DataValidation(type="list", formula1='"EXW,FOB,CIF,DDP,DAP"', allow_blank=True)
    dv_inco.error = 'Please select from dropdown'
    dv_inco.errorTitle = 'Invalid Incoterm'
    ws3.add_data_validation(dv_inco)
    dv_inco.add(f"G2:G1000")  # Column G = INCO1

    sample_data3 = [
        ["0000100001", "1000", "10", "00", "01", "USD", "FOB"],
        ["0000100002", "1000", "10", "00", "02", "GBP", "EXW"],
        ["0000100003", "2000", "10", "00", "01", "EUR", "DDP"],
        ["0000100004", "1000", "20", "00", "03", "JPY", "CIF"],
        ["0000100005", "2000", "10", "00", "01", "EUR", "DAP"]
    ]

    for row_data in sample_data3:
        ws3.append(row_data)

    # Sheet 4: Instructions
    ws4 = wb.create_sheet("Instructions")
    ws4['A1'] = "SAP LSMW Customer Master Data Template - Instructions"
    ws4['A1'].font = Font(bold=True, size=14, color="0070C0")

    instructions = [
        "",
        "How to Use This Template:",
        "1. Fill in customer data in the first three sheets (KNA1, KNB1, KNVV)",
        "2. Use the dropdowns for validated fields (Country, Currency, Incoterms, Account Group)",
        "3. Keep KUNNR consistent across all three sheets for the same customer",
        "4. Review field mapping guide for detailed field descriptions",
        "5. Save as .csv or .txt (tab-delimited) for LSMW upload",
        "",
        "Field Descriptions:",
        "KUNNR - Customer Number (10 chars, use leading zeros)",
        "NAME1 - Customer Name (35 chars max)",
        "LAND1 - Country Key (ISO 2-letter code)",
        "BUKRS - Company Code (4 chars)",
        "AKONT - Reconciliation Account (GL account number)",
        "VKORG - Sales Organization",
        "VTWEG - Distribution Channel",
        "SPART - Division",
        "WAERS - Currency Key",
        "INCO1 - Incoterms",
        "",
        "Data Validation Rules:",
        "- Country codes: US, GB, DE, FR, CN, JP, IN",
        "- Currency: USD, EUR, GBP, JPY, CNY",
        "- Incoterms: EXW, FOB, CIF, DDP, DAP",
        "- Account Groups: KUNA (Sold-to), KUNB (Ship-to), KUNC (Bill-to)",
        "",
        "Important Notes:",
        "- Customer Number (KUNNR) must be 10 characters with leading zeros",
        "- All three tables must have matching KUNNR for the same customer",
        "- Test data should be validated in SAP before production migration",
        "- Use LSMW Step 9 to read this data after saving as tab-delimited text",
        "",
        "For detailed guidance, refer to:",
        "SAP-LSMW-Customer-Master-Migration-Guide.md",
    ]

    for i, text in enumerate(instructions, 1):
        cell = ws4[f'A{i}']
        cell.value = text
        if text.startswith("Field Descriptions:") or text.startswith("Data Validation Rules:") or text.startswith("Important Notes:"):
            cell.font = Font(bold=True, size=12)

    # Adjust column widths for all sheets
    for ws in wb.worksheets:
        for col in ws.columns:
            max_length = 0
            col_letter = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[col_letter].width = min(max_length + 2, 50)

    # Save
    wb.save("Customer-Master-Data-Template.xlsx")
    print("[OK] Excel template created: Customer-Master-Data-Template.xlsx")
    return True


if __name__ == "__main__":
    # Change to deliverables directory
    os.chdir(os.path.dirname(os.path.abspath(__file__)))

    print("=" * 70)
    print("SAP LSMW Document Generator")
    print("=" * 70)
    print()

    success = True

    # Generate Word document
    if not generate_word_document():
        success = False

    print()

    # Generate Excel template
    if not generate_excel_template():
        success = False

    print()
    print("=" * 70)
    if success:
        print("[OK] All documents generated successfully!")
    else:
        print("[ERROR] Some documents failed to generate. Check error messages above.")
    print("=" * 70)

    sys.exit(0 if success else 1)
